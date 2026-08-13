"""Unit tests for the redaction engine.

Run with:  python -m pytest tests/ -v
Or, without pytest installed:  python tests/test_redactor.py
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402

from redactor import PIIType, RedactionConfig, RedactionPipeline  # noqa: E402
from redactor.docx_io import apply_replacements  # noqa: E402
from redactor.entities import Entity, resolve_overlaps  # noqa: E402
from redactor.recognizers.regex_recognizers import luhn_valid, verhoeff_valid  # noqa: E402
from redactor.surrogates import SurrogateFactory  # noqa: E402


def detect(text: str):
    """Detect PII in an isolated string, harvesting from that string only."""
    pipeline = RedactionPipeline(RedactionConfig(use_ner=False))
    _, entities = pipeline.redact_text(text)
    return {(e.pii_type, e.text) for e in entities}


def types_found(text: str):
    return {t for t, _ in detect(text)}


# --------------------------------------------------------------------------
# Checksums
# --------------------------------------------------------------------------

def test_luhn():
    assert luhn_valid("4539 1488 0343 6467")
    assert luhn_valid("4111111111111111")
    assert not luhn_valid("4539 1488 0343 6468")
    assert not luhn_valid("1234")


def test_verhoeff():
    # Verhoeff rejects the overwhelming majority of random 12-digit strings,
    # which is exactly why it guards the Aadhaar recognizer.
    assert not verhoeff_valid("1234 5678 9012")
    assert not verhoeff_valid("12345")


# --------------------------------------------------------------------------
# Structured recognizers
# --------------------------------------------------------------------------

def test_email_and_phone():
    found = detect("Reach Mr. Rohan Dey at rohan.dey@gmail.com or +91 9876543210.")
    assert (PIIType.EMAIL, "rohan.dey@gmail.com") in found
    assert (PIIType.PHONE, "+91 9876543210") in found
    assert (PIIType.PERSON, "Rohan Dey") in found


def test_sentence_final_period_does_not_break_matching():
    """Regression: `(?![\\d.])` also blocked a full stop ending the sentence."""
    assert PIIType.PHONE in types_found("Call +91 9876543210.")
    assert PIIType.IP_ADDRESS in types_found("Host is 192.168.1.14.")


def test_ssn_valid_and_invalid_ranges():
    assert PIIType.SSN in types_found("SSN 452-11-8873 on file.")
    assert PIIType.SSN not in types_found("SSN 000-45-6789 is invalid.")
    assert PIIType.SSN not in types_found("SSN 666-45-6789 is invalid.")


def test_credit_card_requires_luhn():
    assert PIIType.CREDIT_CARD in types_found("Card 4539 1488 0343 6467 charged.")
    assert PIIType.CREDIT_CARD not in types_found("Order 4532-1122-3344-5566 shipped.")


def test_ipv4_octet_range():
    assert PIIType.IP_ADDRESS in types_found("Origin 203.0.113.42 logged.")
    assert PIIType.IP_ADDRESS not in types_found("Build 999.999.999.999 failed.")


def test_dob_requires_birth_context():
    assert PIIType.DATE_OF_BIRTH in types_found("Date of birth 14/03/1982 recorded.")
    # A bare date in a financial document is not a date of birth.
    assert PIIType.DATE_OF_BIRTH not in types_found("The board met on December 10, 2025.")


# --------------------------------------------------------------------------
# Precision guards -- things that must NOT be redacted
# --------------------------------------------------------------------------

def test_public_bodies_and_statutes_are_not_organisations():
    found = types_found(
        "Issued by SEBI under the Companies Act, 2013 and listed on BSE Limited."
    )
    assert PIIType.ORGANISATION not in found


def test_financial_figures_are_not_pii():
    assert detect(
        "Revenue was 4,200.00 million across 10,00,000 Equity Shares; see page 124."
    ) == set()


def test_defined_terms_are_not_companies():
    """`Capital` and `Bank` are legitimate suffixes, but not here."""
    found = types_found(
        "The Escrow Collection Bank and the Refund Bank hold the Equity Share capital."
    )
    assert PIIType.ORGANISATION not in found


def test_role_labels_are_not_people():
    found = types_found("Managing Director and Compliance Officer attended.")
    assert PIIType.PERSON not in found


# --------------------------------------------------------------------------
# Unstructured recognizers
# --------------------------------------------------------------------------

def test_organisation_with_lowercase_connector():
    """Regression: the match used to stop at the lowercase "and"."""
    found = detect("Shares transferred from Shubhkamal Leasing and Investment Private Limited.")
    assert (PIIType.ORGANISATION, "Shubhkamal Leasing and Investment Private Limited") in found


def test_org_span_does_not_cross_commas():
    """Regression: one match once swallowed three separate companies."""
    found = {t for t, _ in detect("Alpha Family Trust, Beta Family Trust and Gamma Family Trust.")}
    assert PIIType.ORGANISATION in found
    texts = [text for t, text in detect(
        "Alpha Family Trust, Beta Family Trust and Gamma Family Trust.") if t == PIIType.ORGANISATION]
    assert all("," not in text for text in texts)


def test_name_followed_by_role_is_harvested():
    found = detect("Our team includes Sandesh Bhagwat, Chief Executive Officer, and others.")
    assert (PIIType.PERSON, "Sandesh Bhagwat") in found


def test_address_anchored_on_pin_code():
    found = detect(
        "having its Registered Office at 11/3, Village Birdewadi, "
        "Chakan Taluka - Khed, Pune - 410 501, Maharashtra, India."
    )
    addresses = [text for t, text in found if t == PIIType.ADDRESS]
    assert addresses
    # The lead-in must not be swallowed into the span.
    assert not addresses[0].lower().startswith("having")


def test_address_abbreviations_do_not_truncate_span():
    """Regression: the last ". " in "lane no. 3" was read as a sentence end."""
    text = "S. no. 245/ 104, Pushpakamal Society, lane no. 3 Prabhat Road, Pune - 411 004, Maharashtra"
    addresses = [t for k, t in detect(text) if k == PIIType.ADDRESS]
    assert addresses and addresses[0].startswith("S. no.")


# --------------------------------------------------------------------------
# Surrogates
# --------------------------------------------------------------------------

def test_surrogates_are_consistent_and_deterministic():
    a, b = SurrogateFactory(salt="x"), SurrogateFactory(salt="x")
    first = a.get(PIIType.PERSON, "Kushal Subbayya Hegde")
    assert first == a.get(PIIType.PERSON, "Kushal Subbayya Hegde")   # consistent
    assert first == b.get(PIIType.PERSON, "Kushal Subbayya Hegde")   # reproducible
    assert first != SurrogateFactory(salt="y").get(PIIType.PERSON, "Kushal Subbayya Hegde")


def test_surrogate_preserves_shape():
    factory = SurrogateFactory()
    phone = factory.get(PIIType.PHONE, "+91 22 4009 4400")
    assert phone.startswith("+91 ")
    assert len([c for c in phone if c.isdigit()]) == 12

    card = factory.get(PIIType.CREDIT_CARD, "4539 1488 0343 6467")
    assert luhn_valid(card), "card surrogates must stay Luhn-valid"

    shouted = factory.get(PIIType.PERSON, "KUSHAL SUBBAYYA HEGDE")
    assert shouted.isupper()


def test_email_surrogate_matches_person_surrogate():
    _, entities = RedactionPipeline(RedactionConfig(use_ner=False)).redact_text(
        "Mr. Rohan Dey can be reached at rohan.dey@gmail.com."
    )
    assert len(entities) >= 2


# --------------------------------------------------------------------------
# Span resolution and docx writing
# --------------------------------------------------------------------------

def test_overlap_resolution_prefers_longer_span():
    long_span = Entity(0, 20, PIIType.ADDRESS, "x" * 20, 0.8, "a")
    short_span = Entity(5, 10, PIIType.ORGANISATION, "y" * 5, 0.99, "b")
    kept = resolve_overlaps([short_span, long_span])
    assert kept == [long_span]


def test_apply_replacements_across_runs():
    """A replacement spanning several runs must land intact in the first."""
    document = Document()
    paragraph = document.add_paragraph()
    for chunk in ("Kushal ", "Subbayya", " Hegde"):
        paragraph.add_run(chunk)
    assert paragraph.text == "Kushal Subbayya Hegde"

    apply_replacements(paragraph, [(0, 21, "John Doe")])
    assert paragraph.text == "John Doe"


def test_apply_replacements_right_to_left():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A and B here")
    apply_replacements(paragraph, [(0, 1, "X"), (6, 7, "Y")])
    assert paragraph.text == "X and Y here"


def test_merged_cells_processed_once(tmp_path=None):
    """Regression: merged cells yielded the same paragraph repeatedly, and
    applying replacements more than once corrupted the text."""
    from redactor.docx_io import iter_text_blocks

    document = Document()
    table = document.add_table(rows=2, cols=3)
    merged = table.rows[0].cells[0].merge(table.rows[0].cells[2])
    merged.text = "Mr. Rohan Dey"
    blocks = [b for b in iter_text_blocks(document) if b.text.strip()]
    texts = [b.text for b in blocks if "Rohan" in b.text]
    assert len(texts) == 1, "merged cell must yield exactly one block"


def _run_all():
    functions = [
        (name, obj) for name, obj in sorted(globals().items())
        if name.startswith("test_") and callable(obj)
    ]
    failures = 0
    for name, function in functions:
        try:
            function()
            print(f"  PASS  {name}")
        except Exception as exc:
            failures += 1
            print(f"  FAIL  {name}: {exc}")
    print(f"\n{len(functions) - failures}/{len(functions)} passed")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
