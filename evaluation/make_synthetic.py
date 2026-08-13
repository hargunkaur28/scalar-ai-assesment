#!/usr/bin/env python3
"""Build a synthetic .docx exercising the PII types the prospectus lacks.

The assignment asks for nine PII types. The supplied Red Herring Prospectus
contains six of them; it has no SSNs, no credit card numbers, no IP addresses
and no dates of birth. Evaluating only against the real document would leave
four of the nine requirements with an empty confusion matrix and no evidence
that the detectors work at all.

This script generates a small document whose gold labels are known by
construction, so those four types get genuine precision and recall numbers.
It also plants deliberate near-misses -- an invalid SSN, a Luhn-failing card
number, an out-of-range IP, dates that are not birthdays, page and order
numbers -- so that precision is measured against the confusions that actually
occur, not just against blank text.

Usage:
    python evaluation/make_synthetic.py --out data/synthetic_pii.docx
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document

#: Each entry is (paragraph text, [(pii_type, surface_text), ...]).
#: An empty entity list marks a NEGATIVE paragraph: everything in it that looks
#: PII-shaped is a decoy that must not be redacted.
CASES: list[tuple[str, list[tuple[str, str]]]] = [
    # ---- Positives -------------------------------------------------------
    (
        "Employee record: Priya Raghunathan, SSN 452-11-8873, date of birth 14/03/1982.",
        [("PERSON", "Priya Raghunathan"), ("SSN", "452-11-8873"),
         ("DATE_OF_BIRTH", "14/03/1982")],
    ),
    (
        "Payment received from card 4539 1488 0343 6467 issued to Marcus Whitfield.",
        [("CREDIT_CARD", "4539 1488 0343 6467"), ("PERSON", "Marcus Whitfield")],
    ),
    (
        "Card on file: 5500-0000-0000-0004. Backup card: 3400 0000 0000 009.",
        [("CREDIT_CARD", "5500-0000-0000-0004"), ("CREDIT_CARD", "3400 0000 0000 009")],
    ),
    (
        "The request originated from 203.0.113.42 and was relayed via 10.0.0.255.",
        [("IP_ADDRESS", "203.0.113.42"), ("IP_ADDRESS", "10.0.0.255")],
    ),
    (
        "IPv6 endpoint 2001:0db8:85a3:0000:0000:8a2e:0370:7334 logged the session.",
        [("IP_ADDRESS", "2001:0db8:85a3:0000:0000:8a2e:0370:7334")],
    ),
    (
        "Applicant Dr. Rohan Dey (DOB: 02 August 1979) can be reached at "
        "rohan.dey@example-mail.com or +91 98765 43210.",
        [("PERSON", "Rohan Dey"), ("DATE_OF_BIRTH", "02 August 1979"),
         ("EMAIL", "rohan.dey@example-mail.com"), ("PHONE", "+91 98765 43210")],
    ),
    (
        "Beneficiary: Ms. Anita Krishnan, born on March 7, 1990, SSN 078-05-1120.",
        [("PERSON", "Anita Krishnan"), ("DATE_OF_BIRTH", "March 7, 1990"),
         ("SSN", "078-05-1120")],
    ),
    (
        "Correspondence address: 48 Riverbend Road, Whitefield, Bengaluru - 560 066, "
        "Karnataka, India.",
        [("ADDRESS", "48 Riverbend Road, Whitefield, Bengaluru - 560 066, Karnataka, India")],
    ),
    (
        "Vendor of record is Northwind Components Private Limited, represented by "
        "Sanjay Iyer.",
        [("ORGANISATION", "Northwind Components Private Limited"), ("PERSON", "Sanjay Iyer")],
    ),
    (
        "PAN AABCK1234M and Aadhaar 2340 1234 5678 were submitted with the form.",
        [("PAN", "AABCK1234M"), ("AADHAAR", "2340 1234 5678")],
    ),

    # ---- Negatives / decoys ---------------------------------------------
    (
        "Order 4532-1122-3344-5566 shipped on 12/01/2024; see invoice 998-12-3456.",
        [],  # Luhn fails, and the 3-2-4 group is an invoice, not an SSN pattern in context
    ),
    (
        "Version 10.0.0.1 of the firmware supersedes build 999.999.999.999.",
        [],  # first is a version string, second is not a valid IPv4
    ),
    (
        "SSN 000-45-6789 and 666-45-6789 are structurally invalid identifiers.",
        [],  # reserved area numbers
    ),
    (
        "The board met on December 10, 2025 and page 124 records the resolution.",
        [],  # a date with no birth context is not a date of birth
    ),
    (
        "Ticket 4111111111111111 was raised against the Companies Act, 2013 by SEBI.",
        [("CREDIT_CARD", "4111111111111111")],
        # Deliberately hard: the number IS Luhn-valid, so it is labelled positive.
        # SEBI and the Companies Act are correctly NOT organisations.
    ),
    (
        "Reserve Bank of India and the National Stock Exchange of India Limited "
        "published the circular on Working Day 3.",
        [],  # allowlisted public bodies
    ),
    (
        "Total revenue was 4,200.00 million against 10,00,000 Equity Shares of "
        "face value 5 each.",
        [],
    ),
    (
        "Refer to the Escrow Collection Bank and the Public Offer Account Bank "
        "as defined above.",
        [],  # defined terms, not company names
    ),
]


def build(out_path: Path) -> Path:
    """Write the synthetic document and its ground truth beside it."""
    document = Document()
    document.add_heading("Synthetic PII Test Document", level=1)
    document.add_paragraph(
        "Generated by evaluation/make_synthetic.py. Every paragraph below has "
        "known gold labels; do not edit without regenerating ground truth."
    )

    blocks = []
    for text, entities in CASES:
        document.add_paragraph(text)
        blocks.append({"text": text, "entities": [list(e) for e in entities]})

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))

    truth_path = out_path.with_suffix(".truth.json")
    truth_path.write_text(
        json.dumps(
            {
                "_comment": "Auto-generated. Regenerate with make_synthetic.py.",
                "source_document": out_path.name,
                "blocks": blocks,
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    return truth_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--out", type=Path, default=Path("data/synthetic_pii.docx"))
    args = parser.parse_args()
    truth = build(args.out)
    print(f"Wrote {args.out}")
    print(f"Wrote {truth}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
